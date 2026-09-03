"""Integration tests for the built-in document adapters."""

from __future__ import annotations

import io
from collections.abc import Iterator

import pytest
from PIL import Image

from document_parser import (
    ContainerBlock,
    DocumentFormat,
    DocumentStatus,
    FigureBlock,
    HeadingBlock,
    ListBlock,
    ParseOptions,
    TableBlock,
    UnsafeDocumentError,
    XlsxOptions,
    convert,
    parse,
)
from document_parser.exceptions import InvalidDocumentError
from document_parser.models import ContentBlock


def walk(blocks: tuple[ContentBlock, ...]) -> Iterator[ContentBlock]:
    for block in blocks:
        yield block
        if isinstance(block, ContainerBlock):
            yield from walk(block.blocks)
        elif isinstance(block, ListBlock):
            for item in block.items:
                yield from walk(item.blocks)
        elif isinstance(block, TableBlock):
            for row in block.rows:
                for cell in row.cells:
                    yield from walk(cell.blocks)


def png_bytes(*, width: int = 32, height: int = 24) -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (width, height), color=(20, 80, 160)).save(output, format="PNG")
    return output.getvalue()


def docx_bytes() -> bytes:
    from docx import Document as WordDocument
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_BREAK

    word = WordDocument()
    word.core_properties.title = "DOCX Sample"
    word.core_properties.author = "Aytan"
    word.sections[0].header.paragraphs[0].text = "Repeated header"
    word.sections[0].footer.paragraphs[0].text = "Repeated footer"
    word.add_heading("Main heading", level=1)
    paragraph = word.add_paragraph()
    paragraph.add_run("Bold").bold = True
    paragraph.add_run(" and italic").italic = True
    word.add_paragraph("First", style="List Bullet")
    word.add_paragraph("Nested", style="List Bullet 2")
    table = word.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "Merged"
    word.add_picture(io.BytesIO(png_bytes()))
    word.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    word.add_section(WD_SECTION.NEW_PAGE)
    word.add_paragraph("Second section")
    output = io.BytesIO()
    word.save(output)
    return output.getvalue()


def xlsx_bytes() -> bytes:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as WorksheetImage
    from openpyxl.worksheet.table import Table

    workbook = Workbook()
    workbook.properties.title = "XLSX Sample"
    worksheet = workbook.active
    worksheet.title = "Visible"
    worksheet.append(["Name", "Amount", "Formula"])
    worksheet.append(["One", 10, "=SUM(B2:B3)"])
    worksheet.append(["Two", 20, None])
    worksheet.add_table(Table(displayName="Data", ref="A1:C3"))
    worksheet.merge_cells("E1:F2")
    worksheet["E1"] = "Merged"
    worksheet.row_dimensions[3].hidden = True
    worksheet.column_dimensions["B"].hidden = True
    worksheet.add_image(WorksheetImage(io.BytesIO(png_bytes())), "H2")
    hidden = workbook.create_sheet("Hidden")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "Secret"
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def pdf_bytes(*, include_scan_page: bool = True) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen.canvas import Canvas

    output = io.BytesIO()
    canvas = Canvas(output, pagesize=letter, pageCompression=0)
    canvas.setTitle("PDF Sample")
    width, height = letter
    canvas.setFont("Helvetica", 10)
    canvas.drawString(40, height - 30, "Repeated header")
    canvas.setFont("Helvetica-Bold", 18)
    canvas.drawString(40, height - 80, "PDF heading")
    canvas.setFont("Helvetica", 11)
    canvas.drawString(40, height - 110, "This is enough native text for digital extraction.")
    for x in (40, 180, 320):
        canvas.line(x, height - 220, x, height - 140)
    for y in (height - 140, height - 180, height - 220):
        canvas.line(40, y, 320, y)
    canvas.drawString(50, height - 165, "Key")
    canvas.drawString(190, height - 165, "Value")
    canvas.drawString(50, height - 205, "A")
    canvas.drawString(190, height - 205, "B")
    canvas.drawString(40, 20, "Repeated footer")
    if include_scan_page:
        canvas.showPage()
        canvas.drawImage(
            ImageReader(io.BytesIO(png_bytes(width=200, height=200))), 0, 0, width, height
        )
    canvas.save()
    return output.getvalue()


def encrypted_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    output = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    writer.write(output)
    return output.getvalue()


def test_docx_adapter_preserves_structure_assets_and_markdown() -> None:
    source = docx_bytes()
    result = convert(source, filename="sample.docx")
    blocks = tuple(walk(result.document.blocks))

    assert result.document.source.format is DocumentFormat.DOCX
    assert result.document.metadata.title == "DOCX Sample"
    assert sum(isinstance(block, ContainerBlock) for block in result.document.blocks) == 2
    assert any(isinstance(block, HeadingBlock) for block in blocks)
    assert any(isinstance(block, ListBlock) for block in blocks)
    assert any(isinstance(block, TableBlock) for block in blocks)
    assert any(isinstance(block, FigureBlock) for block in blocks)
    assert len(result.assets) == 1
    assert "# DOCX Sample" in result.markdown
    assert "Repeated header" not in result.markdown
    assert "sha256-" in result.markdown
    assert convert(source, filename="sample.docx") == result


def test_docx_asset_limits_are_enforced() -> None:
    options = ParseOptions(max_asset_bytes=1, max_total_asset_bytes=1)
    with pytest.raises(UnsafeDocumentError, match="max_asset_bytes"):
        parse(docx_bytes(), filename="large-image.docx", options=options)


def test_xlsx_adapter_preserves_sheets_formulas_merges_and_images() -> None:
    result = convert(xlsx_bytes(), filename="sample.xlsx")
    sheets = tuple(block for block in result.document.blocks if isinstance(block, ContainerBlock))
    blocks = tuple(walk(result.document.blocks))
    tables = tuple(block for block in blocks if isinstance(block, TableBlock))

    assert result.document.source.format is DocumentFormat.XLSX
    assert result.document.metadata.title == "XLSX Sample"
    assert [sheet.attributes["visibility"] for sheet in sheets] == ["visible", "hidden"]
    assert any(
        cell.formula == "=SUM(B2:B3)"
        for table in tables
        for row in table.rows
        for cell in row.cells
    )
    formula_cell = next(
        cell
        for table in tables
        for row in table.rows
        for cell in row.cells
        if cell.formula == "=SUM(B2:B3)"
    )
    assert formula_cell.raw_value == "=SUM(B2:B3)"
    assert formula_cell.number_format == "General"
    assert any(
        cell.column_span == 2 and cell.row_span == 2
        for table in tables
        for row in table.rows
        for cell in row.cells
    )
    assert len(result.assets) == 1
    assert "## Sheet: Visible" in result.markdown
    assert "Sheet: Hidden" not in result.markdown
    assert "=SUM(B2:B3)" in result.markdown


def test_xlsx_options_can_exclude_hidden_content_and_limit_cells() -> None:
    options = ParseOptions(
        xlsx=XlsxOptions(
            include_hidden_sheets=False,
            include_hidden_rows=False,
            include_hidden_columns=False,
            max_worksheet_cells=1_000_000,
        )
    )
    document = parse(xlsx_bytes(), filename="sample.xlsx", options=options)
    assert [
        block.title[0].text for block in document.blocks if isinstance(block, ContainerBlock)
    ] == ["Visible"]

    limited = ParseOptions(xlsx=XlsxOptions(max_worksheet_cells=1))
    with pytest.raises(UnsafeDocumentError, match="max_worksheet_cells"):
        parse(xlsx_bytes(), filename="large.xlsx", options=limited)


def test_pdf_adapter_extracts_native_layout_tables_and_scan_candidates() -> None:
    result = convert(pdf_bytes(), filename="sample.pdf")
    pages = tuple(block for block in result.document.blocks if isinstance(block, ContainerBlock))
    blocks = tuple(walk(result.document.blocks))

    assert result.document.source.format is DocumentFormat.PDF
    assert result.document.metadata.title == "PDF Sample"
    assert result.document.status is DocumentStatus.NEEDS_REVIEW
    assert len(pages) == 2
    assert pages[1].attributes["scan_candidate"] is True
    assert any(diagnostic.code == "pdf.ocr_required" for diagnostic in result.document.diagnostics)
    assert any(isinstance(block, HeadingBlock) for block in blocks)
    assert any(isinstance(block, TableBlock) for block in blocks)
    assert "<!-- page: 1 -->" in result.markdown
    assert result.assets


def test_digital_pdf_is_complete_and_encrypted_pdf_is_rejected() -> None:
    document = parse(pdf_bytes(include_scan_page=False), filename="digital.pdf")
    assert document.status is DocumentStatus.COMPLETE
    with pytest.raises(InvalidDocumentError, match="encrypted PDF"):
        parse(encrypted_pdf_bytes(), filename="encrypted.pdf")
