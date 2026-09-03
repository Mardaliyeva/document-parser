"""Failure isolation and fallback coverage for the native adapters."""

from __future__ import annotations

import hashlib
import io
import zipfile
from datetime import UTC, date, datetime, time, timedelta
from decimal import Decimal
from types import SimpleNamespace
from typing import Any, ClassVar, cast

import pytest

from document_parser import (
    DocumentFormat,
    DocumentStatus,
    DocxOptions,
    ParseOptions,
    PdfOptions,
    UnsafeDocumentError,
    XlsxOptions,
    parse,
)
from document_parser.docx_adapter import (
    DocxAdapter,
    _aware,
    _build_lists,
    _contains_deleted_revision,
    _DocxContext,
    _heading_level,
    _list_spec,
    _ListEntry,
    _ListSpec,
    _note_stories,
    _numbering_map,
    _partition_sections,
    _run_figures,
    _run_span,
    _story_blocks,
)
from document_parser.exceptions import InvalidDocumentError
from document_parser.models import (
    ListKind,
    ParagraphBlock,
    SourceInfo,
    TextSpan,
)
from document_parser.pdf_adapter import (
    PdfAdapter,
    _column_split,
    _find_tables,
    _group_lines,
    _image_boxes,
    _page_blocks,
    _page_images,
    _pdf_date,
    _pdf_metadata,
    _pdf_table,
    _PdfCellData,
    _PdfContext,
    _PdfLine,
    _PdfPageData,
    _PdfTableData,
    _PdfWord,
)
from document_parser.sources import AdapterInput
from document_parser.xlsx_adapter import (
    XlsxAdapter,
    _displayed,
    _occupied_cells,
    _scalar,
    _workbook_metadata,
    _worksheet_images,
    _XlsxContext,
)


def _adapter_input(data: bytes, document_format: DocumentFormat) -> AdapterInput:
    digest = hashlib.sha256(data).hexdigest()
    suffix = document_format.value
    return AdapterInput(
        SourceInfo(
            name=f"sample.{suffix}",
            size_bytes=len(data),
            sha256=digest,
            format=document_format,
            media_type={
                DocumentFormat.DOCX: (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                DocumentFormat.PDF: "application/pdf",
                DocumentFormat.XLSX: (
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ),
            }[document_format],
        ),
        io.BytesIO(data),
    )


def _zip_with(data: bytes, additions: dict[str, bytes]) -> bytes:
    output = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(data)) as source,
        zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target,
    ):
        existing = set(source.namelist())
        for item in source.infolist():
            target.writestr(item, additions.get(item.filename, source.read(item)))
        for name, payload in additions.items():
            if name not in existing:
                target.writestr(name, payload)
    return output.getvalue()


class _FakePdf:
    def __init__(self, pages: list[object]) -> None:
        self.pages = pages

    def __enter__(self) -> _FakePdf:
        return self

    def __exit__(self, *_args: object) -> None:
        return None


class _LayoutPage:
    width = 100
    height = 100
    rotation = 0
    images: ClassVar[tuple[dict[str, object], ...]] = ()

    def extract_words(self, **_kwargs: object) -> list[dict[str, object]]:
        return []

    def find_tables(self, **_kwargs: object) -> list[object]:
        return []


def test_pdf_adapter_isolates_reader_layout_and_page_failures(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pdfplumber
    import pypdf

    data = b"synthetic-pdf"
    source = _adapter_input(data, DocumentFormat.PDF)

    def broken_reader(*_args: object, **_kwargs: object) -> object:
        raise ValueError("bad reader")

    monkeypatch.setattr(pypdf, "PdfReader", broken_reader)
    with pytest.raises(InvalidDocumentError, match="could not be opened"):
        PdfAdapter().parse(source, ParseOptions())

    reader = SimpleNamespace(is_encrypted=False, pages=[], metadata=None)
    monkeypatch.setattr(pypdf, "PdfReader", lambda *_args, **_kwargs: reader)
    monkeypatch.setattr(pdfplumber, "open", lambda *_args, **_kwargs: _FakePdf([_LayoutPage()]))
    mismatch = PdfAdapter().parse(source, ParseOptions(pdf=PdfOptions(detect_tables=False)))
    assert mismatch.document.status is DocumentStatus.PARTIAL
    assert mismatch.document.blocks[0].attributes["scan_candidate"] is False

    class BrokenPage:
        @property
        def width(self) -> int:
            raise ValueError("bad page")

    reader.pages = [SimpleNamespace(images=())]
    monkeypatch.setattr(pdfplumber, "open", lambda *_args, **_kwargs: _FakePdf([BrokenPage()]))
    failed_page = PdfAdapter().parse(source, ParseOptions())
    assert failed_page.document.status is DocumentStatus.PARTIAL
    assert failed_page.document.metadata.title is None
    assert failed_page.document.diagnostics[-1].code == "pdf.page_extraction_failed"

    def broken_layout(*_args: object, **_kwargs: object) -> object:
        raise ValueError("bad layout")

    monkeypatch.setattr(pdfplumber, "open", broken_layout)
    with pytest.raises(InvalidDocumentError, match="layout could not be read"):
        PdfAdapter().parse(source, ParseOptions())


class _Crop:
    def extract_text(self) -> str:
        return "value"


class _RawTable:
    bbox = (0.0, 0.0, 20.0, 30.0)
    cells: ClassVar[tuple[tuple[float, float, float, float], ...]] = (
        (0.0, 0.0, 10.0, 10.0),
        (10.0, 0.0, 20.0, 10.0),
    )
    rows = (1, 2, 3)
    columns = (1, 2)


class _TablePage:
    images: ClassVar[tuple[dict[str, object], ...]] = (
        {"x0": 4, "x1": 4, "top": 2, "bottom": 2},
        {"x0": -1, "x1": 8, "top": -1, "bottom": 8},
    )

    def __init__(self, *, short: bool = False) -> None:
        self.short = short

    def extract_words(self, **_kwargs: object) -> list[dict[str, object]]:
        count = 1 if self.short else 3
        return [
            {
                "text": f"row-{index}",
                "x0": 1,
                "x1": 9,
                "top": index * 10,
                "bottom": index * 10 + 5,
                "size": 10,
            }
            for index in range(count)
        ]

    def find_tables(self, *, table_settings: dict[str, object]) -> list[object]:
        if table_settings["vertical_strategy"] == "lines":
            return []
        rejected = SimpleNamespace(rows=(1,), columns=(1,), cells=(), bbox=(0, 0, 1, 1))
        return [rejected, _RawTable()]

    def crop(self, _box: tuple[float, float, float, float]) -> _Crop:
        return _Crop()


def _page_data(*, tables: tuple[_PdfTableData, ...] = ()) -> _PdfPageData:
    return _PdfPageData(
        page_number=1,
        width=100,
        height=100,
        rotation=0,
        words=(),
        lines=(),
        tables=tables,
        image_boxes=(),
        scan_candidate=False,
    )


def test_pdf_table_and_margin_fallbacks_cover_defensive_paths() -> None:
    assert _image_boxes(_TablePage(), 10, 10) == ((0.0, 0.0, 8.0, 8.0),)
    assert _find_tables(_TablePage(), ParseOptions(pdf=PdfOptions(detect_tables=False))) == ()
    assert _find_tables(_TablePage(short=True), ParseOptions()) == ()
    assert len(_find_tables(_TablePage(), ParseOptions())) == 1

    context = _PdfContext(ParseOptions(), "sample.pdf")
    assert _pdf_table(_PdfTableData((0, 0, 1, 1), ()), _page_data(), context) is None
    invalid = _PdfTableData((0, 0, 1, 1), (_PdfCellData((1, 1, 1, 1), ""),))
    assert _pdf_table(invalid, _page_data(), context) is None

    header = _PdfLine("Header", 1, 20, 1, 5, 20, True)
    footer = _PdfLine("Footer", 1, 20, 95, 99, 10, False)
    page = _page_data(tables=(_PdfTableData((0, 0, 1, 1), ()),))
    page = _PdfPageData(
        page_number=page.page_number,
        width=page.width,
        height=page.height,
        rotation=page.rotation,
        words=page.words,
        lines=(header, footer),
        tables=page.tables,
        image_boxes=page.image_boxes,
        scan_candidate=page.scan_candidate,
    )
    blocks = _page_blocks(page, {"header", "footer"}, 10, [], context, ParseOptions())
    assert [block.attributes["story"] for block in blocks] == ["header", "footer"]

    heading_page = _PdfPageData(
        page_number=1,
        width=100,
        height=100,
        rotation=0,
        words=(),
        lines=(header,),
        tables=(),
        image_boxes=(),
        scan_candidate=False,
    )
    heading = _page_blocks(heading_page, set(), 10, [], context, ParseOptions())[0]
    assert getattr(heading, "level", None) == 6


def test_pdf_two_column_lines_follow_column_reading_order() -> None:
    words = tuple(
        _PdfWord(text, x0, x0 + 30, top, top + 8, 10, "Helvetica")
        for top, left_text, right_text in ((10, "L1", "R1"), (25, "L2", "R2"))
        for text, x0 in ((left_text, 10), (right_text, 200))
    )
    lines = _group_lines(words)
    assert [line.text for line in lines] == ["L1", "R1", "L2", "R2"]
    assert _column_split(list(lines)) == 105

    page = _PdfPageData(
        page_number=1,
        width=300,
        height=100,
        rotation=0,
        words=words,
        lines=lines,
        tables=(),
        image_boxes=(),
        scan_candidate=False,
    )
    blocks = _page_blocks(
        page,
        set(),
        10,
        [],
        _PdfContext(ParseOptions(), "columns.pdf"),
        ParseOptions(pdf=PdfOptions(infer_headings=False)),
    )
    assert [block.spans[0].text for block in blocks if isinstance(block, ParagraphBlock)] == [
        "L1",
        "L2",
        "R1",
        "R2",
    ]

    assert _column_split(lines[:3]) is None
    close_starts = [
        _PdfLine(str(index), index * 10, index * 10 + 5, index * 10, index * 10 + 5, 10, False)
        for index in range(4)
    ]
    assert _column_split(close_starts) is None
    unbalanced = [
        _PdfLine(str(index), x0, x0 + 5, index * 10, index * 10 + 5, 10, False)
        for index, x0 in enumerate((0, 0, 0, 100))
    ]
    assert _column_split(unbalanced) is None


def test_pdf_image_failures_and_limits_are_not_silently_ignored() -> None:
    extracted = _page_data()
    context = _PdfContext(ParseOptions(), "sample.pdf")

    class BrokenImages:
        @property
        def images(self) -> object:
            raise ValueError("bad image list")

    assert _page_images(BrokenImages(), extracted, context) == ()
    assert context.partial is True

    class BrokenImage:
        @property
        def data(self) -> bytes:
            raise ValueError("bad image")

        name = "broken.bin"

    good = SimpleNamespace(data=b"x", name=None)
    context = _PdfContext(ParseOptions(), "sample.pdf")
    figures = _page_images(SimpleNamespace(images=(BrokenImage(), good)), extracted, context)
    assert len(figures) == 1 and figures[0].source is not None
    assert figures[0].source.bounding_box is None
    assert context.partial is True

    limited = _PdfContext(ParseOptions(max_asset_bytes=1, max_total_asset_bytes=1), "sample.pdf")
    with pytest.raises(UnsafeDocumentError, match="max_asset_bytes"):
        _page_images(
            SimpleNamespace(images=(SimpleNamespace(data=b"xx", name="x.bin"),)), extracted, limited
        )

    assert _pdf_date("not-a-date") is None
    assert _pdf_metadata(SimpleNamespace(metadata=None)).title is None


def _workbook_bytes(*, with_chart: bool = False) -> bytes:
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, Reference

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Label", "Value"])
    sheet.append(["A", 1])
    sheet.append(["B", 2])
    if with_chart:
        chart = BarChart()
        chart.add_data(Reference(sheet, min_col=2, min_row=1, max_row=3), titles_from_data=True)
        sheet.add_chart(chart, "D1")
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def test_xlsx_invalid_input_chart_metadata_scalars_and_hidden_merges() -> None:
    with pytest.raises(InvalidDocumentError, match="could not be opened"):
        XlsxAdapter().parse(_adapter_input(b"broken", DocumentFormat.XLSX), ParseOptions())

    chart_document = parse(_workbook_bytes(with_chart=True), filename="chart.xlsx")
    assert chart_document.status is DocumentStatus.PARTIAL
    assert chart_document.diagnostics[-1].code == "xlsx.drawing_omitted"

    properties = SimpleNamespace(
        creator=None,
        keywords=None,
        title=None,
        subject=None,
        created=None,
        modified=None,
        category=None,
    )
    assert _workbook_metadata(SimpleNamespace(properties=properties)).created_at is None
    assert _scalar(Decimal("1.25")) == 1.25
    assert _scalar(date(2026, 1, 2)) == "2026-01-02"
    assert _scalar(datetime(2026, 1, 2, tzinfo=UTC)) == "2026-01-02T00:00:00+00:00"
    assert _scalar(time(3, 4)) == "03:04:00"
    assert _scalar(timedelta(seconds=2)) == "0:00:02"

    class Stable:
        def __str__(self) -> str:
            return "stable"

    assert _scalar(Stable()) == "stable"
    assert _displayed(True) == "TRUE"
    assert _displayed(False) == "FALSE"
    assert _displayed(1.25) == "1.25"

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.merge_cells("A1:B1")
    sheet["A1"] = "merged"
    sheet.column_dimensions["B"].hidden = True
    options = ParseOptions(
        xlsx=XlsxOptions(include_hidden_columns=False, max_worksheet_cells=1_000_000)
    )
    assert _occupied_cells(sheet, options) == {(1, 1)}


def test_xlsx_image_fallbacks_and_asset_limits() -> None:
    class BrokenImage:
        anchor = object()
        format = None
        path = None

        def _data(self) -> bytes:
            raise ValueError("bad image")

    class GoodImage:
        format = None
        path = None

        def __init__(self, anchor: object, data: bytes = b"x") -> None:
            self.anchor = anchor
            self.data = data

        def _data(self) -> bytes:
            return self.data

    worksheet = SimpleNamespace(
        title="Images", _images=[BrokenImage(), GoodImage(object()), GoodImage("C3")]
    )
    context = _XlsxContext(ParseOptions(), "sample.xlsx")
    figures = _worksheet_images(worksheet, context)
    assert len(figures) == 2
    assert figures[0].source is not None and figures[0].source.cell_range is None
    assert figures[1].source is not None and figures[1].source.cell_range == "C3"
    assert context.partial is True

    limited = _XlsxContext(ParseOptions(max_asset_bytes=1, max_total_asset_bytes=1), "sample.xlsx")
    with pytest.raises(UnsafeDocumentError, match="max_asset_bytes"):
        _worksheet_images(
            SimpleNamespace(title="Images", _images=[GoodImage(object(), b"xx")]), limited
        )


def _advanced_docx() -> bytes:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    word = Document()
    paragraph = word.add_paragraph("Visible ")

    relationship_id = paragraph.part.relate_to(
        "https://example.com", RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    for relationship, anchor, text in (
        (relationship_id, None, "External"),
        (None, "bookmark", "Anchor"),
        ("rMissing", None, "Missing"),
    ):
        hyperlink = OxmlElement("w:hyperlink")
        if relationship:
            hyperlink.set(qn("r:id"), relationship)
        if anchor:
            hyperlink.set(qn("w:anchor"), anchor)
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
        if text == "Anchor":
            hyperlink.append(OxmlElement("w:proofErr"))
        paragraph._p.append(hyperlink)

    inserted = OxmlElement("w:ins")
    inserted_run = OxmlElement("w:r")
    inserted_text = OxmlElement("w:t")
    inserted_text.text = "Inserted"
    inserted_run.append(inserted_text)
    inserted.append(inserted_run)
    paragraph._p.append(inserted)

    deleted = OxmlElement("w:del")
    deleted_run = OxmlElement("w:r")
    deleted_text = OxmlElement("w:delText")
    deleted_text.text = "Deleted"
    deleted_run.append(deleted_text)
    deleted.append(deleted_run)
    paragraph._p.append(deleted)

    output = io.BytesIO()
    word.save(output)
    notes = {
        "word/footnotes.xml": (
            b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'<w:footnote w:id="-1"><w:p><w:r><w:t>Separator</w:t></w:r></w:p></w:footnote>'
            b'<w:footnote w:id="1"><w:p><w:r><w:t>Foot note</w:t></w:r></w:p></w:footnote>'
            b'<w:footnote w:id="2"><w:p/></w:footnote></w:footnotes>'
        ),
        "word/endnotes.xml": (
            b'<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'<w:endnote w:id="1"><w:p><w:r><w:t>End note</w:t></w:r></w:p></w:endnote>'
            b"</w:endnotes>"
        ),
    }
    return _zip_with(output.getvalue(), notes)


def test_docx_revisions_hyperlinks_notes_and_story_options() -> None:
    advanced = _advanced_docx()
    document = parse(advanced, filename="advanced.docx")
    texts = [
        span.text
        for block in document.blocks
        for nested in getattr(block, "blocks", ())
        if isinstance(nested, ParagraphBlock)
        for span in nested.spans
    ]
    assert any("Inserted" in text for text in texts)
    assert all("Deleted" not in text for text in texts)
    assert any(item.code == "docx.deleted_revision_omitted" for item in document.diagnostics)
    assert any(
        getattr(block, "attributes", {}).get("story") == "footnotes" for block in document.blocks
    )
    assert any(
        getattr(block, "attributes", {}).get("story") == "endnotes" for block in document.blocks
    )

    options = ParseOptions(
        docx=DocxOptions(
            include_headers=False,
            include_footers=False,
            include_footnotes=False,
            include_endnotes=False,
        )
    )
    without_stories = parse(advanced, filename="advanced.docx", options=options)
    assert all(
        getattr(block, "attributes", {}).get("story") not in {"footnotes", "endnotes"}
        for block in without_stories.blocks
    )

    broken_notes = _zip_with(advanced, {"word/footnotes.xml": b"<broken"})
    with pytest.raises(InvalidDocumentError, match="footnotes"):
        parse(broken_notes, filename="broken-notes.docx")


def test_docx_helper_fallbacks_and_unsupported_drawings(monkeypatch: pytest.MonkeyPatch) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    with pytest.raises(InvalidDocumentError, match="could not be opened"):
        DocxAdapter().parse(_adapter_input(b"broken", DocumentFormat.DOCX), ParseOptions())

    word = Document()
    word.add_section()
    assert len(_partition_sections(word)) == 1
    assert _numbering_map(cast(Any, SimpleNamespace(part=SimpleNamespace()))) == {}
    orphan_number = OxmlElement("w:num")
    orphan_number.set(qn("w:numId"), "999")
    word.part.numbering_part.element.append(orphan_number)
    assert _numbering_map(word)

    paragraph = word.add_paragraph("Numbered")
    properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "2")
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), "7")
    number_properties.extend((level, number_id))
    properties.append(number_properties)
    context = _DocxContext(ParseOptions(), "sample.docx", {(7, 2): (ListKind.ORDERED, 3)})
    assert _list_spec(paragraph, context) == _ListSpec(ListKind.ORDERED, 2, 3)

    outline_paragraph = word.add_paragraph("Outline")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "8")
    outline_paragraph._p.get_or_add_pPr().append(outline)
    assert _heading_level(outline_paragraph) == 6
    hidden = word.add_paragraph().add_run("Hidden")
    hidden.font.hidden = True
    assert _run_span(hidden, None, context) is None
    assert _run_span(
        hidden,
        None,
        _DocxContext(ParseOptions(docx=DocxOptions(include_hidden_text=True)), "x", {}),
    )

    class FakeXml:
        def xpath(self, query: str) -> list[str]:
            if query.endswith("@r:embed"):
                return ["missing"]
            if query.endswith("@r:link"):
                return ["external"]
            return []

        def iter(self) -> list[object]:
            return [
                SimpleNamespace(tag=qn("c:chart")),
                SimpleNamespace(tag=qn("dgm:relIds")),
                SimpleNamespace(tag=qn("m:oMath")),
            ]

    fake_run = SimpleNamespace(_r=FakeXml(), part=SimpleNamespace(related_parts={}))
    assert _run_figures(cast(Any, fake_run), context) == ()
    assert context.partial is True
    assert {item.code for item in context.diagnostics} >= {
        "docx.external_image_omitted",
        "docx.image_missing",
        "docx.chart_omitted",
        "docx.smartart_omitted",
        "docx.equation_simplified",
    }

    entries = (
        _ListEntry(
            _ListSpec(ListKind.UNORDERED, 0, 1),
            (ParagraphBlock(block_id="a", spans=(TextSpan(text="A"),)),),
        ),
        _ListEntry(
            _ListSpec(ListKind.UNORDERED, 1, 1),
            (ParagraphBlock(block_id="b", spans=(TextSpan(text="B"),)),),
        ),
        _ListEntry(
            _ListSpec(ListKind.ORDERED, 0, 4),
            (ParagraphBlock(block_id="c", spans=(TextSpan(text="C"),)),),
        ),
    )
    assert len(_build_lists(entries, context)) == 2
    assert _story_blocks((object(),), context) == ()
    assert _aware(None) is None
    aware = datetime(2026, 1, 1, tzinfo=UTC)
    assert _aware(aware) is aware
    assert _contains_deleted_revision(b"not-a-zip") is False

    valid_data = _advanced_docx()
    monkeypatch.setattr("document_parser.docx_adapter._partition_sections", lambda _word: ((), ()))
    output = DocxAdapter().parse(_adapter_input(valid_data, DocumentFormat.DOCX), ParseOptions())
    assert len(output.document.blocks) >= 2


def test_invalid_note_xml_helper_raises() -> None:
    data = _zip_with(
        _workbook_bytes(),
        {"word/footnotes.xml": b"<broken"},
    )
    context = _DocxContext(ParseOptions(), "sample.docx", {})
    with pytest.raises(InvalidDocumentError, match="footnotes"):
        _note_stories(data, context)

    empty_notes = _zip_with(
        _workbook_bytes(),
        {
            "word/footnotes.xml": (
                b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/'
                b'wordprocessingml/2006/main"><w:footnote w:id="1"><w:p/>'
                b"</w:footnote></w:footnotes>"
            )
        },
    )
    assert _note_stories(empty_notes, context) == ()
